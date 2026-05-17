from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUT_DIR = Path("reports/architecture_review")
DOCX_PATH = OUT_DIR / "mytradingmind_ai_architecture_review.docx"
IMG_OVERVIEW = OUT_DIR / "architecture_overview.png"
IMG_RUNTIME = OUT_DIR / "runtime_flow.png"

BLUE = "2E74B5"
DARK = "0B2545"
MUTED = "6B7280"
LINE = "B8C3CF"
FILL = "F4F7FB"
GOOD = "DFF5EA"
WARN = "FFF4D6"
RISK = "FDE2E2"


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    draw_overview(IMG_OVERVIEW)
    draw_runtime(IMG_RUNTIME)
    build_doc()
    print(DOCX_PATH.resolve())


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], title: str, lines: list[str], fill: str = FILL) -> None:
    draw.rounded_rectangle(xy, radius=18, fill=f"#{fill}", outline=f"#{LINE}", width=3)
    x1, y1, x2, _ = xy
    draw.text((x1 + 24, y1 + 18), title, fill=f"#{DARK}", font=font(30, True))
    y = y1 + 62
    for line in lines:
        draw.text((x1 + 24, y), line, fill="#253244", font=font(22))
        y += 30


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill=f"#{BLUE}", width=4)
    ex, ey = end
    sx, sy = start
    if ex >= sx:
        points = [(ex, ey), (ex - 16, ey - 9), (ex - 16, ey + 9)]
    else:
        points = [(ex, ey), (ex + 16, ey - 9), (ex + 16, ey + 9)]
    draw.polygon(points, fill=f"#{BLUE}")


def draw_overview(path: Path) -> None:
    img = Image.new("RGB", (1800, 1200), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 38), "mytradingmind.ai - Latest Architecture", fill=f"#{DARK}", font=font(44, True))
    draw.text((60, 92), "Binance Spot Testnet live scan, persistent bot instances, hard risk gates, journaling, and validation.", fill=f"#{MUTED}", font=font(24))

    draw_box(draw, (60, 170, 410, 360), "Binance Testnet", ["REST candles", "Websocket trades/books", "Testnet execution"], GOOD)
    draw_box(draw, (500, 150, 880, 390), "Market Data Layer", ["1Y candle features", "Live stream snapshot", "Orderflow metrics"], FILL)
    draw_box(draw, (980, 150, 1360, 390), "Bot Framework", ["Bot instances", "Strategy plugins", "Lifecycle state"], FILL)
    draw_box(draw, (1450, 150, 1760, 390), "Risk Gates", ["Cash/trade", "Risk/trade", "Exposure + kill switch"], RISK)

    draw_box(draw, (60, 500, 410, 720), "Strategy Plugins", ["Existing Momentum", "ATR Trend Burst", "VWAP Reclaim", "Reusable/configurable"], FILL)
    draw_box(draw, (500, 500, 880, 720), "Persistence", ["MariaDB production", "Local JSON fallback", "Parquet/CSV features"], WARN)
    draw_box(draw, (980, 500, 1360, 720), "Journal", ["Bot decisions", "Risk rejections", "Validation results", "Errors/failures"], FILL)
    draw_box(draw, (1450, 500, 1760, 720), "Validation Lab", ["Bot selection", "Period/config", "Metrics + history"], FILL)

    draw_box(draw, (230, 850, 1570, 1085), "Streamlit Operations Dashboard", [
        "Live Trading | Order Flow | Risk | Bot Framework | System Health | Journal | Validation Lab",
        "UI configures, monitors, and reviews; bot/runtime state persists outside the browser session.",
    ], GOOD)

    for s, e in [
        ((410, 265), (500, 265)),
        ((880, 265), (980, 265)),
        ((1360, 265), (1450, 265)),
        ((1170, 390), (1170, 500)),
        ((690, 390), (690, 500)),
        ((1170, 720), (1170, 850)),
        ((690, 720), (690, 850)),
        ((1605, 720), (1470, 850)),
        ((235, 720), (420, 850)),
    ]:
        arrow(draw, s, e)

    img.save(path)


def draw_runtime(path: Path) -> None:
    img = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 38), "Runtime Control and Persistence Flow", fill=f"#{DARK}", font=font(44, True))
    lanes = ["Dashboard", "Bot Framework", "Risk Engine", "Persistence", "Binance Testnet"]
    x_positions = [120, 460, 820, 1180, 1520]
    for x, lane in zip(x_positions, lanes):
        draw.text((x - 70, 130), lane, fill=f"#{DARK}", font=font(28, True))
        draw.line([(x, 180), (x, 950)], fill="#D9E1EA", width=3)

    steps = [
        (120, 250, 460, "create bot"),
        (460, 320, 1180, "persist DRAFT"),
        (120, 400, 460, "deploy bot"),
        (460, 470, 820, "check hard gates"),
        (820, 540, 460, "approved/rejected"),
        (460, 610, 1180, "persist RUNNING/FAILED"),
        (460, 680, 1520, "testnet scan/execution"),
        (1520, 750, 1180, "fills/rejections/latency"),
        (1180, 820, 120, "bot tiles, journal, metrics"),
    ]
    for sx, y, ex, label in steps:
        arrow(draw, (sx, y), (ex, y))
        tx = min(sx, ex) + abs(ex - sx) // 2 - 70
        draw.text((tx, y - 34), label, fill=f"#{BLUE}", font=font(22, True))

    draw_box(draw, (60, 900, 1740, 1015), "Restart Recovery", [
        "On restart, persisted bot state is recovered from MariaDB or local fallback. RUNNING/DEPLOYED bots resume only after risk and protection checks pass."
    ], WARN)
    img.save(path)


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table) -> None:
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
            if row_idx == 0:
                set_cell_fill(cell, "F2F4F7")


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Calibri"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.color.rgb = RGBColor(11, 37, 69)
    for style_name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, "1F4D78")]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    title = doc.add_paragraph()
    title.style = "Title"
    title.add_run("mytradingmind.ai Architecture Review").bold = True
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = subtitle.add_run(f"Bot System, Dashboard, Risk, Journal, and Validation Architecture | Generated {datetime.now():%Y-%m-%d}")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "mytradingmind.ai is now structured around persistent bot instances, reusable strategy plugins, hard portfolio risk gates, Binance Spot Testnet live data, and reviewable bot journaling. "
        "The dashboard controls and monitors the system; bot state, risk settings, journal events, and validation runs persist outside the browser session."
    )

    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Area", "Current Design", "Review Focus"]):
        set_cell_text(cell, text, True)
    rows = [
        ("Market Data", "Binance Spot Testnet websocket and REST candles; one-year Binance feature files for validation.", "Confirm testnet and production endpoint separation."),
        ("Bot Framework", "Strategy-agnostic bot host with DRAFT, BACKTESTED, RUNNING, STOPPED, FAILED lifecycle states.", "Confirm state transitions and recovery rules."),
        ("Risk", "Persisted hard gates for cash/trade, risk/trade, trade window, exposure, and kill switch.", "Confirm gate thresholds before deployment."),
        ("Persistence", "MariaDB schema plus local JSON fallback for desktop mode.", "Confirm MariaDB connection and migration plan."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    style_table(table)

    doc.add_heading("Architecture Overview", level=1)
    doc.add_picture(str(IMG_OVERVIEW), width=Inches(6.8))

    doc.add_heading("Runtime Flow", level=1)
    doc.add_picture(str(IMG_RUNTIME), width=Inches(6.8))

    doc.add_heading("Dashboard Screens", level=1)
    table = doc.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ["Screen", "Purpose", "Key Review Points"]):
        set_cell_text(cell, text, True)
    rows = [
        ("Live Trading", "Market insight dashboard with live ticker, market buckets, and bot instance tiles.", "Tiles should represent each running bot instance clearly."),
        ("Order Flow", "Selected-coin orderflow metrics: pressure, imbalance, spread, velocity, delta, liquidity.", "No generic/global orderflow when a coin is selected."),
        ("Risk", "Portfolio-level capital and exposure controls.", "Risk settings must be hard gates for bot deployment/trading."),
        ("Bot Framework", "Create, configure, backtest, deploy, and stop bot instances.", "Strategies remain pluggable and reusable."),
        ("System Health", "Operational status for websocket, exchange, DB mode, bots, errors, and retries.", "Data refresh should not cause full-page flicker."),
        ("Journal", "Auto-journal of bot lifecycle, decisions, risk blocks, errors, and validation.", "Should answer why a bot entered, exited, skipped, or failed."),
        ("Validation Lab", "Backtest a selected bot over configurable period/capital/fees/slippage.", "Results must persist for later review."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    style_table(table)

    doc.add_heading("Persistence Model", level=1)
    p = doc.add_paragraph()
    p.add_run("Production persistence target: ").bold = True
    p.add_run("MariaDB stores strategies, bot templates, bot instances, runtime state, risk settings, journal events, validation runs, validation metrics, and testnet execution logs. ")
    p.add_run("Local desktop fallback: ").bold = True
    p.add_run("JSON files under reports/ keep the UI operational when MariaDB is disabled.")

    doc.add_heading("Lifecycle and Risk Gate Notes", level=1)
    for item in [
        "Bots run independently from browser sessions; the UI configures and monitors them.",
        "Deployment must pass hard risk gates before a bot can move to RUNNING.",
        "Journal events are created for bot creation, deployment, stopping, risk rejection, and validation.",
        "Validation Lab uses Binance one-year feature data and should expand to direct testnet execution quality metrics.",
        "Restart recovery should resume only persisted DEPLOYED/RUNNING bots after risk and protection checks pass.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Open Review Items", level=1)
    table = doc.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ["Item", "Why It Matters", "Decision Needed"]):
        set_cell_text(cell, text, True)
    rows = [
        ("MariaDB activation", "Desktop fallback works, but 24x7 production requires DB-backed recovery.", "Confirm local MariaDB URL and migration target."),
        ("Worker process", "Long-running bots should execute outside Streamlit.", "Confirm service/process manager for Windows now and Ubuntu later."),
        ("Risk calibration", "Current defaults are conservative placeholders.", "Approve capital, exposure, and trade-window limits."),
        ("Testnet execution logs", "Validation should include fills, rejections, latency, and slippage.", "Confirm order placement mode and credentials."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)
    style_table(table)

    footer = section.footer.paragraphs[0]
    footer.text = "mytradingmind.ai architecture review"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
